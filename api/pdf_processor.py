# api/pdf_processor.py
# DokPDF - Core PDF Processing

import os
import subprocess
import uuid
import shutil
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
import img2pdf
from PIL import Image
import pdf2image

class PDFProcessor:
    """Core PDF processing class for DokPDF"""

    def __init__(self, upload_folder):
        self.upload_folder = upload_folder

    def _get_output_path(self, prefix, ext="pdf"):
        return os.path.join(self.upload_folder, f"{prefix}_{uuid.uuid4().hex}.{ext}")

    def merge_pdfs(self, file_paths):
        """Gabungkan beberapa PDF menjadi satu"""
        try:
            merger = PdfMerger()
            for path in file_paths:
                merger.append(path)

            output_path = self._get_output_path("merged")
            merger.write(output_path)
            merger.close()
            return output_path
        except Exception as e:
            raise Exception(f"Error menggabungkan PDF: {str(e)}")

    def split_pdf(self, file_path, pages=None):
        """Pisahkan PDF berdasarkan halaman"""
        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            writer = PdfWriter()

            if not pages:
                pages = list(range(1, total_pages + 1))

            for page_num in pages:
                if 1 <= page_num <= total_pages:
                    writer.add_page(reader.pages[page_num - 1])

            output_path = self._get_output_path("split")
            writer.write(output_path)
            return output_path
        except Exception as e:
            raise Exception(f"Error memisahkan PDF: {str(e)}")

    def compress_pdf(self, file_path):
        """Kompres PDF"""
        try:
            output_path = self._get_output_path("compressed")
            try:
                cmd = [
                    'gs', '-sDEVICE=pdfwrite', '-dCompatibilityLevel=1.4',
                    '-dPDFSETTINGS=/ebook', '-dNOPAUSE', '-dQUIET', '-dBATCH',
                    f'-sOutputFile={output_path}', file_path
                ]
                subprocess.run(cmd, check=True, capture_output=True)
                return output_path
            except:
                reader = PdfReader(file_path)
                writer = PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)
                writer.write(output_path)
                return output_path
        except Exception as e:
            raise Exception(f"Error mengompres PDF: {str(e)}")

    def pdf_to_images(self, file_path, format_type='png'):
        """Konversi PDF ke gambar"""
        try:
            images = pdf2image.convert_from_path(file_path)
            output_files = []
            for i, image in enumerate(images):
                output_path = self._get_output_path(f"page_{i+1}", format_type)
                image.save(output_path, format_type.upper())
                output_files.append(output_path)
            return output_files
        except Exception as e:
            raise Exception(f"Error konversi PDF ke gambar: {str(e)}")

    def _libreoffice_convert(self, file_path, target_ext):
        output_dir = self.upload_folder
        try:
            cmd = [
                'libreoffice', '--headless', '--convert-to', target_ext,
                '--outdir', output_dir, file_path
            ]
            subprocess.run(cmd, check=True, capture_output=True)
            
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            generated_file = os.path.join(output_dir, f"{base_name}.{target_ext}")
            if os.path.exists(generated_file):
                output_path = self._get_output_path("converted", target_ext)
                shutil.move(generated_file, output_path)
                return output_path
        except:
            pass
        return None

    def pdf_to_word(self, file_path):
        """Konversi PDF ke Word"""
        res = self._libreoffice_convert(file_path, "docx")
        if res: return res
        
        output_path = self._get_output_path("word", "docx")
        from docx import Document
        doc = Document()
        doc.add_heading('DokPDF - Konversi PDF', 0)
        doc.add_paragraph('Gagal konversi penuh. Ini adalah file fallback.')
        doc.save(output_path)
        return output_path

    def pdf_to_excel(self, file_path):
        """Konversi PDF ke Excel"""
        res = self._libreoffice_convert(file_path, "xlsx")
        if res: return res
        raise Exception("Konversi Excel gagal")

    def pdf_to_ppt(self, file_path):
        """Konversi PDF ke PPT"""
        res = self._libreoffice_convert(file_path, "pptx")
        if res: return res
        raise Exception("Konversi PPT gagal")

    def lock_pdf(self, file_path, password):
        """Kunci PDF"""
        try:
            reader = PdfReader(file_path)
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            writer.encrypt(password)
            output_path = self._get_output_path("locked")
            writer.write(output_path)
            return output_path
        except Exception as e:
            raise Exception(f"Error mengunci PDF: {str(e)}")

    def unlock_pdf(self, file_path, password):
        """Buka kunci PDF"""
        try:
            reader = PdfReader(file_path)
            if reader.is_encrypted:
                reader.decrypt(password)
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            output_path = self._get_output_path("unlocked")
            writer.write(output_path)
            return output_path
        except Exception as e:
            raise Exception(f"Error membuka kunci PDF: {str(e)}")

    def add_watermark(self, file_path, text, pages=None):
        """Watermark PDF"""
        try:
            watermark_path = self._get_output_path("wm_temp")
            c = canvas.Canvas(watermark_path, pagesize=letter)
            c.setFont("Helvetica", 40)
            c.setFillGray(0.5, 0.5)
            c.saveState()
            c.translate(300, 400)
            c.rotate(45)
            c.drawCentredString(0, 0, text)
            c.restoreState()
            c.save()

            wm_reader = PdfReader(watermark_path)
            wm_page = wm_reader.pages[0]
            
            reader = PdfReader(file_path)
            writer = PdfWriter()
            
            for i, page in enumerate(reader.pages):
                if not pages or (i + 1) in pages:
                    page.merge_page(wm_page)
                writer.add_page(page)
            
            output_path = self._get_output_path("watermarked")
            writer.write(output_path)
            os.remove(watermark_path)
            return output_path
        except Exception as e:
            raise Exception(f"Error watermark: {str(e)}")

    def organize_pdf(self, file_path, ops):
        """Atur halaman"""
        try:
            reader = PdfReader(file_path)
            writer = PdfWriter()
            order = ops.get('order', list(range(1, len(reader.pages) + 1)))
            for p in order:
                if 1 <= p <= len(reader.pages):
                    writer.add_page(reader.pages[p-1])
            output_path = self._get_output_path("organized")
            writer.write(output_path)
            return output_path
        except Exception as e:
            raise Exception(f"Error organize: {str(e)}")

    def add_signature(self, pdf_path, sig_path, position):
        """Tanda tangan"""
        try:
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            
            sig_img = Image.open(sig_path)
            temp_sig_pdf = self._get_output_path("sig_temp")
            c = canvas.Canvas(temp_sig_pdf, pagesize=letter)
            
            # Simple position logic
            x, y = 450, 50
            if position == 'top-left': x, y = 50, 700
            elif position == 'top-right': x, y = 450, 700
            elif position == 'bottom-left': x, y = 50, 50
            
            c.drawImage(sig_path, x, y, width=100, height=50, mask='auto')
            c.save()
            
            sig_reader = PdfReader(temp_sig_pdf)
            sig_page = sig_reader.pages[0]
            
            last_page = reader.pages[-1]
            last_page.merge_page(sig_page)
            
            for page in reader.pages:
                writer.add_page(page)
                
            output_path = self._get_output_path("signed")
            writer.write(output_path)
            os.remove(temp_sig_pdf)
            return output_path
        except Exception as e:
            raise Exception(f"Error signature: {str(e)}")